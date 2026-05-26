#!/usr/bin/env python3
"""
Korean Legal Opinion Markdown → DOCX converter.

Implements the DOCX defaults in legal-writing-formatting-guide.md §17.1:
- A4 (210mm x 297mm), 2.54cm margins
- Times New Roman 11pt (Latin) + 맑은 고딕 11pt (CJK), eastAsia explicitly set
- Body line spacing 1.15, paragraph after 6pt
- Statute blocks: bordered single-cell tables, line spacing 1.0, after 3pt
- Black text, no italic, bold for emphasis only

Usage:
    python3 scripts/md-to-docx.py <input.md> <output.docx>

Designed for legal-writing-agent output (opinion.md). Handles:
- ATX headings (#, ##, ###, ####, #####)
- Blockquotes (>) → bordered statute blocks
- Markdown tables (|...|...|) → docx tables (info block, risk matrix, citation grade)
- Inline **bold**, *italic*, `code`
- Numbered/lettered enumeration as paragraphs (text preserved verbatim)
- --- horizontal rules (skipped)
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ASCII_FONT = "Times New Roman"
CJK_FONT = "맑은 고딕"
BODY_SIZE_PT = 11
HEADING_SIZE = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11.5}


def _resolve_private_dir(project_root: Path) -> Path:
    return Path(
        os.environ.get("LEGAL_ORCHESTRATOR_PRIVATE_DIR", str(project_root / "output"))
    ).expanduser()


def _resolve_work_product_path(raw_arg: str, project_root: Path) -> Path:
    path = Path(raw_arg).expanduser()
    if path.is_absolute():
        return path.resolve()
    if path.parts and path.parts[0] in {"output", "samples", ".", ".."}:
        return (project_root / path).resolve()

    private_candidate = _resolve_private_dir(project_root) / path
    if len(path.parts) >= 2 and private_candidate.parent.exists():
        return private_candidate.resolve()

    return (project_root / path).resolve()


def set_run_font(run, size_pt: float = BODY_SIZE_PT, bold: bool = False, italic: bool = False, mono: bool = False) -> None:
    """Apply Times New Roman + 맑은 고딕 (eastAsia) to a run."""
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    ascii_font = "Courier New" if mono else ASCII_FONT
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)
    rFonts.set(qn("w:cs"), ascii_font)


INLINE_PATTERN = re.compile(
    r"(\*\*[^*\n]+?\*\*|`[^`\n]+?`|\*[^*\s][^*\n]*?\*)"
)
_ESCAPE_TAG_RE = re.compile(r"<escape>(.*?)</escape>", re.DOTALL)
ESCAPED_OMISSION_TEXT = "[Sanitized instruction-like text omitted]"


def _render_escape_tags(md: str, *, preserve_escaped_text: bool = False) -> str:
    """Render sanitizer wrappers according to the delivery policy."""
    replacement = r"\1" if preserve_escaped_text else ESCAPED_OMISSION_TEXT
    return _ESCAPE_TAG_RE.sub(replacement, md)


def add_inline_runs(para, text: str, base_bold: bool = False, size_pt: float = BODY_SIZE_PT) -> None:
    """Tokenize inline markdown and emit runs with proper font."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            set_run_font(run, size_pt=size_pt, bold=base_bold)
        token = m.group()
        if token.startswith("**"):
            run = para.add_run(token[2:-2])
            set_run_font(run, size_pt=size_pt, bold=True)
        elif token.startswith("`"):
            run = para.add_run(token[1:-1])
            set_run_font(run, size_pt=size_pt, bold=base_bold, mono=True)
        elif token.startswith("*"):
            run = para.add_run(token[1:-1])
            set_run_font(run, size_pt=size_pt, bold=base_bold, italic=True)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        set_run_font(run, size_pt=size_pt, bold=base_bold)


def style_paragraph(para, line_spacing: float = 1.15, space_after_pt: float = 6, space_before_pt: float = 0) -> None:
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    space_before = {1: 18, 2: 14, 3: 10, 4: 8, 5: 6}.get(level, 6)
    space_after = {1: 10, 2: 8, 3: 6, 4: 4, 5: 4}.get(level, 4)
    style_paragraph(p, line_spacing=1.15, space_after_pt=space_after, space_before_pt=space_before)
    add_inline_runs(p, text, base_bold=True, size_pt=HEADING_SIZE.get(level, 11))


def _set_borders(cell, sz: str = "4") -> None:
    """Apply 0.5pt single-line borders to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, twips: int = 100) -> None:
    """Set cell internal margins (twips; 1pt = 20 twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(twips))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def add_statute_block(doc, lines: list[str]) -> None:
    """Render a `> ...` blockquote as a single-cell bordered table (statute box)."""
    cleaned: list[str] = []
    for ln in lines:
        if ln.startswith("> "):
            cleaned.append(ln[2:])
        elif ln.startswith(">"):
            cleaned.append(ln[1:].lstrip(" "))
        else:
            cleaned.append(ln)
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    if not cleaned:
        return

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_borders(cell)
    _set_cell_margins(cell, twips=140)

    first = True
    for ln in cleaned:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        style_paragraph(p, line_spacing=1.0, space_after_pt=3, space_before_pt=0)
        if ln.strip():
            add_inline_runs(p, ln, size_pt=BODY_SIZE_PT)

    spacer = doc.add_paragraph()
    style_paragraph(spacer, space_after_pt=6)


def parse_table_rows(buf: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in buf:
        ln = raw.strip()
        if not ln:
            continue
        if re.match(r"^\|[\s\-:|]+\|$", ln):
            continue
        if ln.startswith("|") and ln.endswith("|"):
            cells = [c.strip() for c in ln[1:-1].split("|")]
            rows.append(cells)
    return rows


def add_md_table(doc, buf: list[str]) -> None:
    rows = parse_table_rows(buf)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            _set_borders(cell)
            _set_cell_margins(cell, twips=80)
            p = cell.paragraphs[0]
            style_paragraph(p, line_spacing=1.0, space_after_pt=2, space_before_pt=0)
            add_inline_runs(p, cell_text, base_bold=(i == 0), size_pt=BODY_SIZE_PT - 0.5)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, space_after_pt=6)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = ASCII_FONT
    style.font.size = Pt(BODY_SIZE_PT)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ASCII_FONT)
    rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)
    rFonts.set(qn("w:cs"), ASCII_FONT)


def convert(md_path: Path, docx_path: Path, *, preserve_escaped_text: bool = False) -> None:
    text = _render_escape_tags(
        md_path.read_text(encoding="utf-8"),
        preserve_escaped_text=preserve_escaped_text,
    )
    lines = text.split("\n")

    doc = Document()
    setup_document(doc)

    quote_buf: list[str] = []
    table_buf: list[str] = []
    i = 0
    while i < len(lines):
        ln = lines[i]

        # Blockquote accumulation
        if ln.startswith(">"):
            if table_buf:
                add_md_table(doc, table_buf)
                table_buf = []
            quote_buf.append(ln)
            i += 1
            continue
        if quote_buf:
            add_statute_block(doc, quote_buf)
            quote_buf = []

        # Table accumulation
        stripped = ln.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(ln)
            i += 1
            continue
        if table_buf:
            add_md_table(doc, table_buf)
            table_buf = []

        # Horizontal rule — skip
        if stripped == "---":
            i += 1
            continue

        # Headings
        m = re.match(r"^(#+)\s+(.+)$", ln)
        if m:
            level = min(len(m.group(1)), 5)
            add_heading(doc, m.group(2).strip(), level=level)
            i += 1
            continue

        # Common opinion titles (centered, large, bold)
        if stripped in {
            "**MEMORANDUM**",
            "**법률 검토 의견서**",
            "**Legal Analysis Memo**",
        }:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            style_paragraph(p, space_after_pt=14, space_before_pt=6)
            r = p.add_run(stripped.strip("*"))
            set_run_font(r, size_pt=20, bold=True)
            i += 1
            continue

        # Date line directly under the opinion title (e.g. "2026. 4. 10.") — center it
        if re.match(r"^\d{4}\.\s*\d{1,2}\.\s*\d{1,2}\.$", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            style_paragraph(p, space_after_pt=12, space_before_pt=0)
            r = p.add_run(stripped)
            set_run_font(r, size_pt=BODY_SIZE_PT)
            i += 1
            continue

        # Empty line — paragraph break (skip; spacing handled by previous para)
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        style_paragraph(p, line_spacing=1.15, space_after_pt=6)
        add_inline_runs(p, ln)
        i += 1

    # Flush trailing buffers
    if quote_buf:
        add_statute_block(doc, quote_buf)
    if table_buf:
        add_md_table(doc, table_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Convert legal-opinion Markdown to DOCX.")
    parser.add_argument("input_md")
    parser.add_argument("output_docx")
    parser.add_argument(
        "--preserve-escaped-text",
        action="store_true",
        help="Preserve text inside <escape> tags instead of omitting it from the rendered DOCX.",
    )
    args = parser.parse_args(argv)
    project_root = Path.cwd().resolve()
    md = _resolve_work_product_path(args.input_md, project_root)
    docx = _resolve_work_product_path(args.output_docx, project_root)
    if not md.exists():
        print(f"Input not found: {md}", file=sys.stderr)
        return 1
    convert(md, docx, preserve_escaped_text=args.preserve_escaped_text)
    size = docx.stat().st_size
    print(f"Saved {docx} ({size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
